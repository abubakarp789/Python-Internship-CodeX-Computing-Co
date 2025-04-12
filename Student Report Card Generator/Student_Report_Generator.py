import os
import json
from typing import Dict, List, Optional
from dataclasses import dataclass
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER

# Constants
REPORTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'reports')
DATA_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'student_data.json')
os.makedirs(REPORTS_DIR, exist_ok=True)

@dataclass
class Student:
    name: str
    roll_number: str
    subjects: Dict[str, float]
    total_marks: float = 0.0
    average: float = 0.0
    grade: str = ""

    def add_subject(self, subject: str, marks: float) -> bool:
        try:
            marks = float(marks)
            if not 0 <= marks <= 100:
                raise ValueError("Marks must be between 0 and 100")
            self.subjects[subject] = marks
            return True
        except ValueError as e:
            print(f"Error: {str(e)}")
            return False

    def calculate_results(self) -> bool:
        if not self.subjects:
            return False
            
        self.total_marks = sum(self.subjects.values())
        self.average = self.total_marks / len(self.subjects)
        
        # Determine grade based on average marks
        if self.average >= 90:
            self.grade = "A+"
        elif self.average >= 80:
            self.grade = "A"
        elif self.average >= 70:
            self.grade = "B"
        elif self.average >= 60:
            self.grade = "C"
        elif self.average >= 50:
            self.grade = "D"
        else:
            self.grade = "F"
        
        return True

    def to_dict(self) -> dict:
        return {
            'name': self.name,
            'roll_number': self.roll_number,
            'subjects': self.subjects,
            'total_marks': self.total_marks,
            'average': self.average,
            'grade': self.grade
        }

    @classmethod
    def from_dict(cls, data: dict) -> 'Student':
        student = cls(
            name=data['name'],
            roll_number=data['roll_number'],
            subjects=data['subjects']
        )
        student.total_marks = data['total_marks']
        student.average = data['average']
        student.grade = data['grade']
        return student

class ReportCardGenerator:
    def __init__(self):
        self.students: List[Student] = []
        self.load_data()

    def load_data(self) -> None:
        try:
            if os.path.exists(DATA_FILE):
                with open(DATA_FILE, 'r') as f:
                    data = json.load(f)
                    self.students = [Student.from_dict(student_data) for student_data in data]
        except Exception as e:
            print(f"Error loading data: {e}")

    def save_data(self) -> None:
        try:
            with open(DATA_FILE, 'w') as f:
                json.dump([student.to_dict() for student in self.students], f, indent=4)
        except Exception as e:
            print(f"Error saving data: {e}")

    def add_student(self, name: str, roll_number: str) -> Optional[Student]:
        if self.find_student(roll_number):
            print(f"Error: Student with roll number {roll_number} already exists.")
            return None
            
        student = Student(name, roll_number, {})
        self.students.append(student)
        self.save_data()
        return student

    def find_student(self, roll_number: str) -> Optional[Student]:
        for student in self.students:
            if student.roll_number == roll_number:
                return student
        return None

    def generate_text_report(self, student: Student) -> Optional[str]:
        if not student.subjects:
            print("No subjects added for this student.")
            return None
            
        student.calculate_results()
        
        report = []
        report.append("=" * 60)
        report.append(f"STUDENT REPORT CARD".center(60))
        report.append("=" * 60)
        report.append(f"Name: {student.name}")
        report.append(f"Roll Number: {student.roll_number}")
        report.append("-" * 60)
        report.append(f"{'Subject':<30}{'Marks':>10}")
        report.append("-" * 60)
        
        for subject, marks in student.subjects.items():
            report.append(f"{subject:<30}{marks:>10.2f}")
        
        report.append("-" * 60)
        report.append(f"{'Total Marks':<30}{student.total_marks:>10.2f}")
        report.append(f"{'Average':<30}{student.average:>10.2f}")
        report.append(f"{'Grade':<30}{student.grade:>10}")
        report.append("=" * 60)
        
        return "\n".join(report)

    def save_report_to_text_file(self, student: Student) -> bool:
        report_text = self.generate_text_report(student)
        if not report_text:
            return False
            
        try:
            filename = os.path.join(REPORTS_DIR, f"{student.name}_{student.roll_number}_report.txt")
            with open(filename, 'w') as file:
                file.write(report_text)
            print(f"Report card saved as '{filename}'")
            return True
        except Exception as e:
            print(f"Error saving text file: {e}")
            return False

    def save_report_to_word(self, student: Student) -> bool:
        if not student.subjects:
            print("No subjects added for this student.")
            return False
            
        student.calculate_results()
        
        try:
            doc = Document()
            
            # Add title
            doc.add_heading('STUDENT REPORT CARD', 0)
            
            # Add student details
            doc.add_paragraph(f'Name: {student.name}')
            doc.add_paragraph(f'Roll Number: {student.roll_number}')
            
            # Add subject and marks table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Subject'
            header_cells[1].text = 'Marks'
            
            # Add subjects and marks
            for subject, marks in student.subjects.items():
                row_cells = table.add_row().cells
                row_cells[0].text = subject
                row_cells[1].text = f"{marks:.2f}"
            
            # Add summary row
            doc.add_paragraph('')
            summary_table = doc.add_table(rows=3, cols=2)
            summary_table.style = 'Table Grid'
            
            # Add total, average, and grade
            total_row = summary_table.rows[0].cells
            total_row[0].text = 'Total Marks'
            total_row[1].text = f"{student.total_marks:.2f}"
            
            avg_row = summary_table.rows[1].cells
            avg_row[0].text = 'Average'
            avg_row[1].text = f"{student.average:.2f}"
            
            grade_row = summary_table.rows[2].cells
            grade_row[0].text = 'Grade'
            grade_row[1].text = student.grade
            
            # Save document
            filename = os.path.join(REPORTS_DIR, f"{student.name}_{student.roll_number}_report.docx")
            doc.save(filename)
            print(f"Report card saved as Word document: '{filename}'")
            return True
        except Exception as e:
            print(f"Error saving Word document: {e}")
            return False

    def save_report_to_pdf(self, student: Student) -> bool:
        if not student.subjects:
            print("No subjects added for this student.")
            return False
            
        student.calculate_results()
        
        try:
            filename = os.path.join(REPORTS_DIR, f"{student.name}_{student.roll_number}_report.pdf")
            doc = SimpleDocTemplate(filename, pagesize=letter)
            
            # Create content list
            elements = []
            
            # Add styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Heading1'],
                alignment=TA_CENTER,
                spaceAfter=12
            )
            
            # Add title
            elements.append(Paragraph("STUDENT REPORT CARD", title_style))
            elements.append(Spacer(1, 12))
            
            # Add student details
            elements.append(Paragraph(f"Name: {student.name}", styles["Normal"]))
            elements.append(Paragraph(f"Roll Number: {student.roll_number}", styles["Normal"]))
            elements.append(Spacer(1, 12))
            
            # Create subject and marks table
            subject_data = [['Subject', 'Marks']]
            for subject, marks in student.subjects.items():
                subject_data.append([subject, f"{marks:.2f}"])
            
            subject_table = Table(subject_data, colWidths=[300, 100])
            subject_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            elements.append(subject_table)
            elements.append(Spacer(1, 12))
            
            # Create summary table
            summary_data = [
                ['Total Marks', f"{student.total_marks:.2f}"],
                ['Average', f"{student.average:.2f}"],
                ['Grade', student.grade]
            ]
            
            summary_table = Table(summary_data, colWidths=[300, 100])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            elements.append(summary_table)
            
            # Build PDF
            doc.build(elements)
            print(f"Report card saved as PDF document: '{filename}'")
            return True
        except Exception as e:
            print(f"Error saving PDF document: {e}")
            return False

class UserInterface:
    def __init__(self):
        self.generator = ReportCardGenerator()

    def get_valid_input(self, prompt: str, validation_func) -> str:
        while True:
            try:
                value = input(prompt).strip()
                if validation_func(value):
                    return value
            except Exception as e:
                print(f"Error: {str(e)}")

    def validate_name(self, name: str) -> bool:
        if not name:
            raise ValueError("Name cannot be empty")
        if not name.replace(' ', '').isalpha():
            raise ValueError("Name should contain only letters and spaces")
        return True

    def validate_roll_number(self, roll_number: str) -> bool:
        if not roll_number:
            raise ValueError("Roll number cannot be empty")
        if not roll_number.isalnum():
            raise ValueError("Roll number should contain only letters and numbers")
        return True

    def validate_marks(self, marks: str) -> bool:
        try:
            marks_float = float(marks)
            if not 0 <= marks_float <= 100:
                raise ValueError("Marks must be between 0 and 100")
            return True
        except ValueError:
            raise ValueError("Please enter a valid number for marks")

    def add_student(self) -> None:
        print("\n=== Add New Student ===")
        name = self.get_valid_input("Enter student name: ", self.validate_name)
        roll_number = self.get_valid_input("Enter roll number: ", self.validate_roll_number)
        
        student = self.generator.add_student(name, roll_number)
        if student:
            print(f"Student '{name}' added successfully!")

    def add_subjects(self) -> None:
        print("\n=== Add Subjects and Marks ===")
        roll_number = self.get_valid_input("Enter student roll number: ", self.validate_roll_number)
        student = self.generator.find_student(roll_number)
        
        if not student:
            print("Student not found.")
            return
            
        print(f"\nAdding subjects for {student.name}")
        while True:
            subject = input("\nEnter subject name (or 'done' to finish): ").strip()
            if subject.lower() == 'done':
                break
                
            if not subject:
                print("Subject name cannot be empty")
                continue
                
            marks = self.get_valid_input(f"Enter marks for {subject}: ", self.validate_marks)
            if student.add_subject(subject, marks):
                print(f"Subject '{subject}' added successfully!")
                self.generator.save_data()

    def view_report(self) -> None:
        print("\n=== View Report Card ===")
        roll_number = self.get_valid_input("Enter student roll number: ", self.validate_roll_number)
        student = self.generator.find_student(roll_number)
        
        if not student:
            print("Student not found.")
            return
            
        if not student.subjects:
            print("No subjects added for this student yet.")
            return
            
        print("\nStudent Report Card:")
        report_text = self.generator.generate_text_report(student)
        print(report_text)

    def save_report(self) -> None:
        print("\n=== Save Report Card ===")
        roll_number = self.get_valid_input("Enter student roll number: ", self.validate_roll_number)
        student = self.generator.find_student(roll_number)
        
        if not student:
            print("Student not found.")
            return
            
        if not student.subjects:
            print("No subjects added for this student yet.")
            return
            
        print("\nSave report card as:")
        print("1. Text file")
        print("2. Word document")
        print("3. PDF document")
        print("4. All formats")
        
        while True:
            choice = input("Enter your choice (1-4): ").strip()
            if choice in ['1', '2', '3', '4']:
                break
            print("Invalid choice. Please enter 1-4.")
        
        if choice == '1':
            self.generator.save_report_to_text_file(student)
        elif choice == '2':
            self.generator.save_report_to_word(student)
        elif choice == '3':
            self.generator.save_report_to_pdf(student)
        elif choice == '4':
            self.generator.save_report_to_text_file(student)
            self.generator.save_report_to_word(student)
            self.generator.save_report_to_pdf(student)

    def run(self) -> None:
        print("STUDENT REPORT CARD GENERATOR")
        print("=" * 30)
        
        while True:
            print("\nMenu:")
            print("1. Add a new student")
            print("2. Add subjects and marks")
            print("3. View report card")
            print("4. Save report card")
            print("5. Exit")
            
            choice = input("\nEnter your choice (1-5): ").strip()
            
            if choice == '1':
                self.add_student()
            elif choice == '2':
                self.add_subjects()
            elif choice == '3':
                self.view_report()
            elif choice == '4':
                self.save_report()
            elif choice == '5':
                print("Thank you for using the Student Report Card Generator!")
                break
            else:
                print("Invalid choice. Please try again.")

if __name__ == "__main__":
    ui = UserInterface()
    ui.run()